import os
import re
from docx import Document

# Ruta del archivo limpio
archivo_limpio = r"D:\LIBROS\Cómo ganar amigos e influir sobre las personas ( PDFDrive )\Libro_Limpio.docx"

# Cargar el documento
doc = Document(archivo_limpio)

# Ruta de la carpeta donde se guardarán los archivos
carpeta_destino = os.path.dirname(archivo_limpio)

# Variables de control
secciones = {}
current_section = "Prefacio"
secciones[current_section] = []

# Expresión regular mejorada para detectar títulos de capítulos en cualquier formato
regex_titulo = re.compile(r"^\s*(PREFACIO|PARTE\s*\d+|CAP[IÍ]TULO\s*\d+|SECCI[ÓO]N\s*\d+|EP[IÍ]LOGO|INTRODUCCI[ÓO]N)", re.IGNORECASE)

# Dividir el documento por capítulos
for para in doc.paragraphs:
    texto = para.text.strip()

    # Si el texto coincide con un título de capítulo/parte, crear nueva sección
    if regex_titulo.match(texto):
        current_section = re.sub(r"[^a-zA-Z0-9áéíóúÁÉÍÓÚ ]", "", texto)  # Limpiar caracteres raros
        secciones[current_section] = []

    # Agregar el contenido a la sección actual
    secciones[current_section].append(texto)

# Guardar cada sección en un archivo separado
for idx, (titulo, contenido) in enumerate(secciones.items()):
    if "prefacio" in titulo.lower():
        nombre_archivo = f"Prefacio_Cómo ganar amigos e influir sobre las personas ( PDFDrive ).docx"
    elif "parte" in titulo.lower():
        nombre_archivo = f"Parte_{idx}_Cómo ganar amigos e influir sobre las personas ( PDFDrive ).docx"
    elif "capítulo" in titulo.lower():
        numero = re.search(r"\d+", titulo)
        nombre_archivo = f"Capitulo_{numero.group()}_Cómo ganar amigos e influir sobre las personas ( PDFDrive ).docx" if numero else f"Capitulo_{idx}_Cómo ganar amigos e influir sobre las personas ( PDFDrive ).docx"
    else:
        nombre_archivo = f"Seccion_{idx}_Cómo ganar amigos e influir sobre las personas ( PDFDrive ).docx"

    ruta_guardado = os.path.join(carpeta_destino, nombre_archivo)

    # Crear un nuevo documento y guardar el contenido
    nuevo_doc = Document()
    nuevo_doc.add_paragraph(titulo)  
    nuevo_doc.add_paragraph("\n".join(contenido))  
    nuevo_doc.save(ruta_guardado)

    print(f"✅ Archivo creado: {ruta_guardado}")

print("\n✅ Todas las secciones han sido guardadas correctamente.")
