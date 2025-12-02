from docx import Document
from docx.shared import Pt

# Crear documento
doc = Document()

# Función para agregar párrafos con tamaño de fuente más compacto
def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    if bold:
        run.bold = True

# Agregar nombre y datos de contacto
doc.add_heading('Arnold Torres Larriega', level=1)
add_paragraph("CIP: 290066")
add_paragraph("Teléfono: (+51) 944 206 911  |  Correo: arnoldtorresla@gmail.com  |  LinkedIn: linkedin.com/in/arnoldtorrespmp")

# Sección de resumen
doc.add_heading('RESUMEN PROFESIONAL', level=2)
add_paragraph(
    "Ingeniero Industrial con más de 5 años de experiencia en gestión de proyectos y 2 años con certificación PMP. "
    "Especializado en telecomunicaciones, logística, desarrollo de software y consultoría, liderando hasta 80 personas "
    "y gestionando proyectos de alto impacto. Amplio dominio de Scrum, Kanban y PMBOK 7ª edición."
)
add_paragraph(
    "Habilidad para la optimización de procesos, análisis de datos y toma de decisiones estratégicas. "
    "Desarrollo en Python, automatizando tareas repetitivas de gestión, logrando una reducción del 99.97% en tiempo total de procesamiento "
    "(de 24 horas-hombre a menos de 5 minutos). Capacidad para identificar oportunidades de mejora e implementar soluciones tecnológicas innovadoras "
    "que generan impacto positivo."
)

# Sección de experiencia
doc.add_heading('EXPERIENCIA PROFESIONAL', level=2)

experiencias = [
    ("Project Manager | MSI Americas", "Oct 2023 - Actualidad", [
        "Gestión de 6 proyectos simultáneos en telecomunicaciones, logística, pruebas de señal de campo y desarrollo de software.",
        "Implementación de metodologías ágiles e híbridas, adaptadas a cada cliente.",
        "Dirección de 80 personas, asegurando cumplimiento de cronogramas y objetivos.",
        "Automatización en Python, reduciendo el tiempo de procesamiento en 99.97%.",
        "Negociación y gestión de stakeholders con Huawei, Nokia, Claro y Entel, incluyendo interacción en inglés con clientes de empresas transnacionales.",
        "Gestión de proyectos con ingresos superiores a $1.5M USD.",
        "Cumplimiento de los altos estándares y exigencias de clientes internacionales, incluyendo empresas del mercado asiático.",
        "Experiencia en Drive Test, asegurando la calidad de red en entornos urbanos y rurales."
    ]),
    ("Project Manager | A&M Ingeniería Y Proyectos S.A.C", "Feb 2022 - Oct 2023", [
        "Implementación de Scrum y Kanban para optimización de procesos.",
        "Mitigación de riesgos operativos y financieros, asegurando cumplimiento de hitos críticos."
    ]),
    ("Project Manager | JA & DE Ingenieros S.A", "Feb 2022 - Oct 2023", [
        "Gestión de proyectos de ingeniería conforme a estándares internacionales.",
        "Uso de análisis de datos para toma de decisiones estratégicas."
    ])
]

for puesto, periodo, detalles in experiencias:
    add_paragraph(puesto, bold=True)
    add_paragraph(periodo)
    for item in detalles:
        add_paragraph(f"- {item}")

# Sección de educación y certificaciones
doc.add_heading('EDUCACIÓN Y CERTIFICACIONES', level=2)
add_paragraph("Ingeniería Industrial | Universidad Ricardo Palma (2013 - 2019)")
add_paragraph("Project Management Professional (PMP) | PMI (2023 - Actualidad)")

# Habilidades y conocimientos
doc.add_heading('HABILIDADES Y CONOCIMIENTOS', level=2)
skills = [
    "Gestión de Proyectos: PMP, Scrum, Kanban, PMBOK 7ª edición.",
    "Optimización de Procesos: Mejora continua, gestión de riesgos.",
    "Análisis de Datos: Business Intelligence, toma de decisiones basadas en datos.",
    "Automatización en Python: Reducción del 99.97% del tiempo de procesamiento.",
    "Gestión de Telecomunicaciones: Implementación de proyectos y pruebas de señal de campo (Drive Test).",
    "Herramientas Tecnológicas: SAP, MS Project, Jira, Trello, SharePoint, Power BI, CRM.",
    "Inteligencia Artificial: Generación de prompts avanzados para IA y desarrollo asistido.",
    "Negociación con clientes internacionales, asegurando cumplimiento de altos estándares de calidad."
]
for skill in skills:
    add_paragraph(f"- {skill}")

# Valor agregado a la empresa
doc.add_heading('VALOR AGREGADO A LA EMPRESA', level=2)
valores = [
    "Optimización de procesos mediante metodologías ágiles y herramientas digitales.",
    "Análisis de datos y Business Intelligence para toma de decisiones estratégicas.",
    "Automatización de tareas con IA y Python, mejorando eficiencia y rentabilidad.",
    "Liderazgo en ejecución de proyectos, asegurando resultados tangibles y escalables.",
    "Adaptabilidad en entornos dinámicos, con enfoque en transformación digital."
]
for valor in valores:
    add_paragraph(f"- {valor}")

# Guardar el documento en la ruta especificada
ruta = r"D:\AR DOCS\Arnold Torres cv 2025.docx"
doc.save(ruta)

print(f"El CV ha sido generado exitosamente en {ruta}")
