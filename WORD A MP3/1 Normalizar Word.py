from docx import Document

# Ruta del documento original
archivo_original = r"D:\LIBROS\Cómo ganar amigos e influir sobre las personas ( PDFDrive )\Cómo ganar amigos e influir sobre las personas ( PDFDrive ).docx"

# Cargar el documento
doc = Document(archivo_original)

# Crear un nuevo documento limpio
nuevo_doc = Document()

# Limpiar el texto (eliminar espacios extra, líneas vacías, etc.)
for para in doc.paragraphs:
    texto = para.text.strip()
    if texto:  # Evitar líneas vacías
        nuevo_doc.add_paragraph(texto)

# Guardar el documento limpio
archivo_limpio = r"D:\LIBROS\Cómo ganar amigos e influir sobre las personas ( PDFDrive )\Libro_Limpio.docx"
nuevo_doc.save(archivo_limpio)

print(f"✅ Documento limpio guardado en: {archivo_limpio}")
